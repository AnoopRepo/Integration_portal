"""
document_reader.py — Converts binary document files into plain text strings.

WHY THIS FILE EXISTS:
  This is the bridge between the physical world (files on disk) and
  the language model world (text strings).

  Qwen is a text model. It cannot open files. It cannot read binary data.
  This module does the reading so Qwen only ever sees plain text.

  WHAT EACH LIBRARY DOES:
    PyMuPDF (fitz):
      - Industry-standard PDF text extractor
      - Reads PDF page-by-page and extracts text layer
      - Works on most searchable PDFs (not scanned image PDFs)

    python-docx (docx.Document):
      - Reads Microsoft Word .docx files
      - Extracts text from paragraphs
      - Does NOT read text in tables or headers by default
      (table reading can be added in Phase 2)

    Built-in open():
      - For .txt files — fastest and simplest

WHY TUPLE RETURN (text, error) NOT EXCEPTIONS:
  When processing 10 resumes, if 1 fails due to OneDrive sync,
  we want to continue with the other 9 — not crash the whole pipeline.
  Returning (None, error_message) lets the caller handle it gracefully.

DATA FLOW:
  file path (string)
       ↓ detect extension (.pdf / .docx / .txt)
       ↓ call appropriate reader function
  (text_string, None)    ← on success
  (None, error_message)  ← on failure
       ↓ returned to knowledge_agent.py
"""

import os
from pathlib import Path
from typing import Tuple, Optional


def read_pdf(path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract all text from a PDF file using PyMuPDF.

    HOW IT WORKS:
      fitz.open() loads the PDF into memory.
      We loop through every page and call page.get_text().
      get_text() extracts the text layer embedded in the PDF.
      NOTE: Scanned/image PDFs have no text layer — they return empty strings.
            A future Phase could add OCR (pytesseract) for those.

    Input:  "C:/Users/.../Anoop-J-Resume.pdf"
    Output: ("Anoop J\nEmail: anoop@gmail.com\nSkills: Java...", None)
         or (None, "File is locked by OneDrive sync. Try again.")

    The seven error types handled:
      ImportError     — pymupdf not installed
      PermissionError — file locked by OneDrive/antivirus
      FileNotFoundError — file is cloud-only (not downloaded)
      fitz.FileDataError — corrupted or password-protected PDF
      Exception       — catch-all for unexpected errors
    """
    try:
        import fitz  # PyMuPDF — installed as 'pymupdf' package
    except ImportError:
        return None, "PyMuPDF is not installed. Run: pip install pymupdf"

    try:
        doc = fitz.open(path)
        text_parts = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)

        doc.close()
        full_text = "\n".join(text_parts).strip()

        if not full_text:
            return None, (
                "PDF has no extractable text. It may be a scanned image PDF. "
                "Only text-based PDFs are supported in Phase 1."
            )

        return full_text, None

    except PermissionError:
        return None, (
            "File is locked — possibly syncing with OneDrive. "
            "Try again in a moment, or set the folder to 'Always keep on this device'."
        )
    except FileNotFoundError:
        return None, (
            "File not found on disk. It may be stored in the cloud only (OneDrive). "
            "Set the folder to 'Always keep on this device' in OneDrive settings."
        )
    except Exception as e:
        err_type = type(e).__name__
        return None, f"PDF read failed ({err_type}): {str(e)}"


def read_docx(path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract all text from a Word .docx file using python-docx.

    HOW IT WORKS:
      Document(path) loads the .docx XML structure.
      doc.paragraphs is a list of Paragraph objects.
      Each paragraph has a .text property — the plain text string.
      We filter empty paragraphs and join with newlines.

    NOTE: python-docx reads paragraph text only.
          Text inside tables, headers, footers is NOT included by default.
          This covers 90% of resume formats. Table reading = Phase 2.

    Input:  "C:/Users/.../Ashish_Deloitte (2).pdf"   ← path with spaces/parens, handled by pathlib
    Output: ("Ashish Kumar\nJava Developer\nSkills: Java...", None)
         or (None, "File is locked. Try again.")
    """
    try:
        from docx import Document  # python-docx package
    except ImportError:
        return None, "python-docx is not installed. Run: pip install python-docx"

    try:
        doc = Document(path)

        # Extract paragraph text, skip empty paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables (bonus: handles table-format resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() and cell.text.strip() not in paragraphs:
                        paragraphs.append(cell.text.strip())

        full_text = "\n".join(paragraphs).strip()

        if not full_text:
            return None, "DOCX appears empty — no text paragraphs or table cells found."

        return full_text, None

    except PermissionError:
        return None, "File is locked — possibly syncing with OneDrive. Try again in a moment."
    except FileNotFoundError:
        return None, "File not found on disk. It may be cloud-only in OneDrive."
    except Exception as e:
        err_type = type(e).__name__
        return None, f"DOCX read failed ({err_type}): {str(e)}"


def read_txt(path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Read a plain .txt file.

    WHY errors="replace":
      Some text files use non-UTF-8 encodings (old Windows files use CP1252).
      errors="replace" substitutes unreadable characters with '?' instead of crashing.

    Input:  "C:/Users/.../resume.txt"
    Output: ("John Doe\nSkills: Python, SQL...", None)
    """
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read().strip()

        if not text:
            return None, "TXT file is empty."

        return text, None

    except PermissionError:
        return None, "File is locked. Try again in a moment."
    except FileNotFoundError:
        return None, "File not found on disk."
    except Exception as e:
        return None, f"TXT read failed ({type(e).__name__}): {str(e)}"


def read_csv(path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract all text from a CSV file.

    HOW IT WORKS:
      Reads the CSV row-by-row and joins elements with a '|' separator.
      This preserves row structure so downstream LLM prompts or metadata
      extractors can make sense of tabular data.

    Input:  "C:/Users/.../billing.csv"
    Output: ("Item | Quantity | Price\nBook | 1 | 10.99...", None)
    """
    try:
        import csv
        text_lines = []
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            for row in reader:
                if row:
                    # Filter out empty elements and join with pipe
                    cleaned_row = [cell.strip() for cell in row]
                    text_lines.append(" | ".join(cleaned_row))
        
        full_text = "\n".join(text_lines).strip()
        if not full_text:
            return None, "CSV file is empty."
        
        return full_text, None

    except PermissionError:
        return None, "File is locked — possibly syncing with OneDrive. Try again in a moment."
    except FileNotFoundError:
        return None, "File not found on disk."
    except Exception as e:
        return None, f"CSV read failed ({type(e).__name__}): {str(e)}"


_FILE_CONTENT_CACHE = {}  # maps absolute path -> (mtime, size, text, error)


def read_document(path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Master dispatch function — detects file type and calls the right reader.
    Caches parsed text in memory to avoid redundant file system reading.
    The cache automatically invalidates if the file modification time or size changes.

    WHY THIS PATTERN (dispatch function):
      The caller (knowledge_agent.py) doesn't need to know which reader
      to use. It just calls read_document(path) for any file type.
      This is the "facade pattern" — one clean entry point hides
      the implementation details.

    Input:  Any file path string
    Output: (text, None) on success
            (None, error_message) on failure

    Supported: .pdf → read_pdf()  |  .docx → read_docx()  |  .txt → read_txt()  |  .csv → read_csv()
    """
    if path.startswith("knowledge_layer://"):
        record_id = path.replace("knowledge_layer://", "")
        from app.knowledge.repositories.stub_repository import knowledge_repo
        import asyncio
        try:
            try:
                loop = asyncio.get_event_loop()
            except RuntimeError:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
            record = loop.run_until_complete(knowledge_repo.get(record_id))
            if record:
                return record.content, None
            else:
                return None, f"Knowledge record '{record_id}' not found"
        except Exception as e:
            return None, f"Failed to retrieve knowledge record: {e}"

    abs_path = os.path.abspath(path)
    
    # Check modification time and file size for caching
    try:
        mtime = os.path.getmtime(abs_path)
        size = os.path.getsize(abs_path)
        
        if abs_path in _FILE_CONTENT_CACHE:
            c_mtime, c_size, cached_text, cached_err = _FILE_CONTENT_CACHE[abs_path]
            if c_mtime == mtime and c_size == size:
                # Cache hit!
                return cached_text, cached_err
    except Exception:
        # If stat fails, bypass cache checks
        mtime = None
        size = None

    file_path = Path(abs_path)
    suffix = file_path.suffix.lower()
    filename = file_path.name

    print(f"[Reader] Reading '{filename}' ({suffix})")

    if suffix == ".pdf":
        text, error = read_pdf(abs_path)
    elif suffix == ".docx":
        text, error = read_docx(abs_path)
    elif suffix in (".txt", ".py", ".json"):
        text, error = read_txt(abs_path)
    elif suffix == ".csv":
        text, error = read_csv(abs_path)
    else:
        return None, f"Unsupported file type '{suffix}'. Supported: .pdf, .docx, .txt, .csv, .py, .json"

    if text:
        print(f"[Reader] '{filename}' -> {len(text)} characters extracted")
    else:
        print(f"[Reader] '{filename}' -> ERROR: {error}")

    # Cache the result if we got stat details
    if mtime is not None and size is not None:
        try:
            _FILE_CONTENT_CACHE[abs_path] = (mtime, size, text, error)
        except Exception:
            pass

    return text, error
